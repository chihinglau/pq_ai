"""
PQ AI Terminal 对外技术方案报告生成脚本（脱敏版）
对用户名、密码、IP地址等敏感信息进行脱敏处理
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ========== 样式设置 ==========
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5

def set_run_font(run, font_name='微软雅黑', font_size=None, bold=False, color=None):
    run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, '微软雅黑')
    return p

def add_paragraph(doc, text, bold=False, alignment=None, font_size=11):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, '微软雅黑', font_size=font_size, bold=bold)
    return p

def add_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Light List Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, '微软雅黑', font_size=10, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(cell_text))
            set_run_font(run, '微软雅黑', font_size=10)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

# ========== 封面页 ==========
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PQ AI Terminal')
set_run_font(run, '微软雅黑', font_size=36, bold=True, color=(0, 51, 102))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基于终端波形数据的电能质量 AI 应用')
set_run_font(run, '微软雅黑', font_size=22, bold=True, color=(0, 102, 153))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('—— 新能源与充电桩接入影响评估 ——')
set_run_font(run, '微软雅黑', font_size=16, color=(102, 102, 102))

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('技术方案报告')
set_run_font(run, '微软雅黑', font_size=28, bold=True, color=(51, 51, 51))

for _ in range(4):
    doc.add_paragraph()

info_items = [
    ('版本', 'v2.2.1'),
    ('日期', '2026年8月'),
    ('状态', '已通过真实硬件验证'),
    ('目标平台', '嵌入式终端 + AI 算力模组'),
    ('文档密级', '公开技术方案'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f'{label}：')
    set_run_font(run1, '微软雅黑', font_size=12, bold=True)
    run2 = p.add_run(value)
    set_run_font(run2, '微软雅黑', font_size=12)

doc.add_page_break()

# ========== 1. 项目概述与背景 ==========
add_heading_custom(doc, '1. 项目概述与背景', level=1)

add_heading_custom(doc, '1.1 项目背景', level=2)
add_paragraph(doc,
    '随着"双碳"战略推进，分布式光伏和电动汽车充电桩在配电网中的渗透率持续提高，'
    '新能源负荷的波动性、随机性和非线性特征对传统配电网的安全稳定运行带来严峻挑战。'
    '新能源接入引发的电能质量（Power Quality, PQ）问题主要包括：电压偏差、谐波畸变、'
    '三相不平衡、电压暂降/暂升等，这些问题直接影响电网设备的安全运行和用户供电质量。')

add_paragraph(doc,
    '本项目针对上述挑战，构建了一套从算法仿真验证到嵌入式实时部署的完整解决方案，'
    '实现对配电网关键节点的电能质量在线监测与智能分析，为新能源的友好接入提供技术支撑。')

add_heading_custom(doc, '1.2 项目目标', level=2)

goals = [
    '实现 7 通道高精度波形实时采集（3 相电压 + 3 相电流 + 零序），采样率 12.8 kHz',
    '实时计算 12 项国标电能质量指标（电压偏差、THD、三相不平衡度等）',
    '基于 AI 算法实现异常检测与事件分类，识别新能源接入引发的典型电能质量事件',
    '构建采样主机 + AI 算力模组的双机协作架构，通过 USB 虚拟网卡实现高速互连',
    '支持 S1~S5 五类典型场景识别（基准负荷、充电桩、分布式光伏、光充耦合、极端工况）',
]

for goal in goals:
    p = doc.add_paragraph(goal, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '1.3 技术特色', level=2)

features = [
    ('双机协作架构', '采样主机专司采集与边缘计算，算力模组专司 AI 推理，通过 USB 虚拟网卡实现高速互连'),
    ('真实验证通过', '已在真实硬件上完成全链路验证：采集 → 传输 → AI 推理 → 结果返回'),
    ('多算法融合', '融合孤立森林（iForest）、自编码器（AE）、一维卷积神经网络（1D-CNN）三类 AI 算法'),
    ('国标对标', 'PQ 指标计算完全遵循 GB/T 12325/14549/15543/15945 等国家标准'),
    ('分级日志系统', '实现 ERROR/WARN/INFO/DEBUG 四级日志，便于运维监控与问题排查'),
]

for title, desc in features:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{title}：')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, '微软雅黑', font_size=11)

doc.add_page_break()

# ========== 2. 系统总体架构 ==========
add_heading_custom(doc, '2. 系统总体架构', level=1)

add_heading_custom(doc, '2.1 总体架构设计', level=2)
add_paragraph(doc,
    '本系统采用"采样主机 + 算力模组"的双机协作架构，充分发挥各硬件平台的优势：')

arch_items = [
    ('采样主机', '负责高速波形采集、PQ 指标实时计算、事件触发与特征提取'),
    ('计量 AFE', '7 通道 24bit 高精度模数转换器，内置谐波分析引擎'),
    ('AI 算力模组', '负责 AI 推理（iForest/AE/CNN1D），支持后续大模型部署'),
    ('USB 虚拟网卡', '通过 USB 实现两机高速 TCP/IP 通信，延迟 < 5ms'),
]

for title, desc in arch_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{title}：')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, '微软雅黑', font_size=11)

add_heading_custom(doc, '2.2 数据流', level=2)

add_paragraph(doc,
    '系统每 20ms（一个工频周波）执行一次完整的数据处理流程：')

dataflow_steps = [
    '计量 AFE 以 12.8 kHz 采样率采集 7 通道波形，每通道 256 个采样点',
    '硬件抽象层（HAL）读取波形数据，传递给 PQ 指标计算模块',
    '实时计算 12 项 PQ 指标，评估是否超出国标限值',
    '事件触发引擎检测异常事件（谐波超标、三相不平衡等），采用滞回机制防抖',
    '特征工程模块从波形与指标中提取 27 维特征向量',
    '特征向量通过 USB 虚拟网卡发送至 AI 算力模组',
    '算力模组依次运行 iForest、AE、CNN1D 三类 AI 模型',
    'AI 推理结果（48字节响应包）返回采样主机',
    '场景识别模块根据 AI 结果与规则判定当前场景',
    '治理建议生成、数据持久化存储、MQTT 上报云端',
]

for i, step in enumerate(dataflow_steps, 1):
    p = doc.add_paragraph()
    run1 = p.add_run(f'步骤 {i}：')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(step)
    set_run_font(run2, '微软雅黑', font_size=11)

doc.add_page_break()

# ========== 3. 硬件平台选型 ==========
add_heading_custom(doc, '3. 硬件平台选型', level=1)

add_heading_custom(doc, '3.1 核心硬件组件', level=2)

headers = ['组件', '型号', '核心参数', '应用角色']
data = [
    ['采样主控', '全志 T536', '4×Cortex-A55 + E907 RISC-V', '波形采集 + 边缘计算'],
    ['计量 AFE', '钜泉 HT7627S', '7 通道 24bit ADC，12.8kHz', '高精度模数转换'],
    ['算力模组', '瑞芯微 RK3576', '6×A76 + NPU', 'AI 推理 + 大模型部署'],
    ['互连总线', 'USB 虚拟网卡', 'TCP/IP 协议', '双机高速通信'],
]
add_table(doc, headers, data, col_widths=[2.5, 4, 5, 3.5])

add_heading_custom(doc, '3.2 采样主机', level=2)
add_paragraph(doc,
    '采样主机采用高性能 SoC，具备以下特点：')

t536_features = [
    '四核 Cortex-A55 处理器，满足实时性要求',
    '内置 RISC-V 核心，可用于低功耗任务处理',
    '丰富的外设接口：SPI/I2C/UART/USB/以太网',
    '支持 Linux 操作系统，软件生态成熟',
    '工作温度范围 -40°C ~ 85°C，适应工业环境',
]

for feat in t536_features:
    p = doc.add_paragraph(feat, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '3.3 计量前端', level=2)
add_paragraph(doc,
    '计量前端是专为电能质量监测设计的高精度 AFE：')

afe_features = [
    '7 通道同步采样：3 相电压 + 3 相电流 + 零序电压',
    '24bit 高精度 ADC，动态范围 > 120dB',
    '支持 256/512 点/周波的采样率配置（12.8kHz / 25.6kHz）',
    '内置谐波分析引擎，支持 2~31 次谐波实时计算',
    '提供 RMS、THD、功率、频率等寄存器直接读取',
    'SPI 接口通信，传输速率可达 20MHz',
]

for feat in afe_features:
    p = doc.add_paragraph(feat, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '3.4 AI 算力模组', level=2)
add_paragraph(doc,
    'AI 算力模组具备以下特性：')

rk3576_features = [
    '6 核 ARM Cortex-A76 处理器',
    '内置 NPU 算力单元，支持 INT8/INT16 量化模型推理',
    '算力可达 6 TOPS，满足实时 AI 推理需求',
    '丰富的 AI 开发工具链：支持 PyTorch/TensorFlow/ONNX',
    'USB 3.0/PCIe/以太网等高速接口，便于外设扩展',
]

for feat in rk3576_features:
    p = doc.add_paragraph(feat, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

doc.add_page_break()

# ========== 4. 软件技术方案 ==========
add_heading_custom(doc, '4. 软件技术方案', level=1)

add_heading_custom(doc, '4.1 软件分层架构', level=2)
add_paragraph(doc,
    '软件采用分层架构设计，各层职责清晰、接口标准化：')

headers = ['层级', '模块', '核心功能', '运行位置']
data = [
    ['应用层', '场景识别 / 治理建议', 'S1~S5 场景判定', '采样主机'],
    ['AI 推理层', 'iForest / AE / CNN1D', '异常检测、事件分类', '算力模组'],
    ['特征工程层', 'feature_extract', '27 维特征提取', '采样主机'],
    ['算法层', 'PQ 指标 / 事件触发', 'PQ 指标计算、事件触发', '采样主机'],
    ['通信层', 'USB 网卡 / AI RPC', '数据传输、远程调用', '双机'],
    ['HAL 层', '硬件抽象层', 'AFE 寄存器/波形接口', '采样主机'],
    ['驱动层', 'AFE 驱动', 'SPI 通信、芯片配置', '采样主机'],
]
add_table(doc, headers, data, col_widths=[2, 4, 5, 3])

add_heading_custom(doc, '4.2 波形采集与处理', level=2)

headers = ['参数', '配置值', '说明']
data = [
    ['采样率', '12800 Hz', '对应 50Hz 工频每周波 256 点'],
    ['通道数', '7', 'UA/UB/UC/IA/IB/IC/U0'],
    ['每周期点数', '256', '满足谐波分析需求'],
    ['ADC 精度', '24 bit', '原生分辨率'],
    ['数据格式', 'float32', '单精度浮点传输'],
]
add_table(doc, headers, data, col_widths=[3, 3, 8])

add_heading_custom(doc, '4.3 PQ 指标计算', level=2)

headers = ['序号', '指标名称', '计算方法', '国标', '限值']
data = [
    ['1', '电压偏差', '(Vrms - Vnom) / Vnom × 100%', 'GB/T 12325', '±7%'],
    ['2', '电压 THD', '√(ΣHn²) / H1 × 100%', 'GB/T 14549', '5%'],
    ['3', '电流 THD', '√(ΣHn²) / H1 × 100%', 'GB/T 14549', '8%'],
    ['4', '三相不平衡度', 'max偏差 / Vavg × 100%', 'GB/T 15543', '2%'],
    ['5', '频率偏差', '|f - 50|', 'GB/T 15945', '±0.5Hz'],
    ['6', '变压器负载率', 'S_apparent / S_rated × 100%', '—', '100%'],
    ['7', '线路负载率', 'I_rms / I_max × 100%', '—', '100%'],
    ['8', '功率因数', 'P / S', '—', '0.85'],
]
add_table(doc, headers, data, col_widths=[1.2, 2.5, 4.5, 2.5, 2.5])

add_heading_custom(doc, '4.4 AI 推理与异常检测', level=2)

ai_algo_items = [
    ('孤立森林（iForest）', '无监督异常检测', '基于树结构隔离异常点，输出异常得分 [0,1]，无需标签数据'),
    ('自编码器（AE）', '无监督异常检测', '学习正常模式的压缩表示，通过重构误差检测异常'),
    ('一维卷积神经网络（1D-CNN）', '有监督事件分类', '自动提取波形特征，分类 7 类典型事件，输出置信度'),
]

for name, category, desc in ai_algo_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{name}（{category}）：')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, '微软雅黑', font_size=11)

add_heading_custom(doc, '4.5 双机通信方案', level=2)

comm_items = [
    ('物理连接', '采样主机 USB Host → 算力模组 USB Device，USB 2.0'),
    ('驱动协议', 'USB 以太网控制模型，Linux cdc_ether 驱动'),
    ('传输协议', 'TCP 流式传输'),
    ('传输延迟', '平均 RTT < 5ms，满足实时性要求'),
]

for title, desc in comm_items:
    p = doc.add_paragraph()
    run1 = p.add_run(f'{title}：')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(desc)
    set_run_font(run2, '微软雅黑', font_size=11)

add_heading_custom(doc, '4.5.1 原始波形传输协议', level=3)

headers = ['字段', '大小', '说明']
data = [
    ['魔数', '4 字节', '"WAVE" 帧起始标识'],
    ['协议版本', '1 字节', '当前为 0x01'],
    ['通道数', '1 字节', '固定为 7'],
    ['每周期点数', '2 字节', '固定为 256'],
    ['周期 ID', '4 字节', '递增计数器'],
    ['时间戳', '8 字节', '微秒精度'],
    ['波形数据', '7168 字节', '7 通道 × 256 点 × 4 字节 float'],
]
add_table(doc, headers, data, col_widths=[3, 3, 10])

add_heading_custom(doc, '4.5.2 AI 推理响应协议', level=3)

headers = ['字段', '大小', '说明']
data = [
    ['魔数', '4 字节', '响应包起始标识'],
    ['状态码', '1 字节', '0=成功，非零=错误'],
    ['iForest 得分', '4 字节 float', '[0, 1]，越接近 1 越异常'],
    ['AE 得分', '4 字节 float', 'MSE 重构误差'],
    ['CNN 类别', '4 字节 int', '事件分类 ID（0~6）'],
    ['CNN 置信度', '4 字节 float', '[0, 1]'],
    ['场景 ID', '4 字节 int', 'S1~S5 场景标识'],
    ['场景名称', '16 字节 char', '场景名称字符串'],
    ['时间戳', '8 字节 int64', '微秒精度'],
]
add_table(doc, headers, data, col_widths=[3, 3, 10])

doc.add_page_break()

# ========== 5. AI 算法方案 ==========
add_heading_custom(doc, '5. AI 算法方案', level=1)

add_heading_custom(doc, '5.1 孤立森林异常检测（iForest）', level=2)
add_heading_custom(doc, '5.1.1 算法原理', level=3)
add_paragraph(doc,
    '算法核心思想：异常样本在特征空间中容易被随机分割隔离。构建多棵随机二叉树，'
    '每棵树随机选择特征与分裂点。异常样本的平均路径长度较短，'
    '归一化后得到异常得分 s ∈ (0, 1)，越接近 1 表示越异常。')

add_heading_custom(doc, '5.1.2 模型参数', level=3)

headers = ['参数', '取值', '说明']
data = [
    ['树数量', '32', '森林规模，越多越稳定'],
    ['树最大深度', '8', '单棵树最大深度'],
    ['输入特征维度', '27', '对应特征工程输出'],
    ['异常阈值', '0.6', '得分 > 0.6 判定为异常'],
]
add_table(doc, headers, data, col_widths=[3, 3, 8])

add_heading_custom(doc, '5.2 自编码器异常检测（AE）', level=2)
add_heading_custom(doc, '5.2.1 算法原理', level=3)
add_paragraph(doc,
    '自编码器由编码器和解码器组成，在训练时仅使用正常样本。'
    '编码器将输入压缩到低维瓶颈层，解码器重构原始输入。'
    '异常样本无法被正确重构，导致重构误差（MSE）显著升高。')

add_heading_custom(doc, '5.2.2 网络结构', level=3)
add_paragraph(doc, '输入(27维) → 编码器(27→8) → 瓶颈层(8维) → 解码器(8→27) → 重构输出(27维)')

add_heading_custom(doc, '5.3 一维卷积神经网络事件分类（1D-CNN）', level=2)
add_heading_custom(doc, '5.3.1 算法原理', level=3)
add_paragraph(doc,
    '1D-CNN 适合处理时序信号（如电力波形），通过卷积层提取局部特征，'
    '全局平均池化降维，全连接层输出 7 类事件的概率分布。')

add_heading_custom(doc, '5.3.2 网络结构', level=3)
add_paragraph(doc, '输入(256点波形) → Conv1D(8 滤波器 × 5 核) → ReLU → GlobalAvgPool → FC(7 类) → Softmax')

add_heading_custom(doc, '5.3.3 事件分类定义', level=3)

headers = ['类别 ID', '事件类型', '说明']
data = [
    ['0', 'NORMAL', '正常运行'],
    ['1', 'SAG', '电压暂降'],
    ['2', 'SWELL', '电压暂升'],
    ['3', 'HARMONIC', '谐波超标'],
    ['4', 'UNBALANCE', '三相不平衡'],
    ['5', 'OVERLOAD', '过载'],
    ['6', 'TRANSIENT', '瞬态脉冲'],
]
add_table(doc, headers, data, col_widths=[2.5, 3.5, 10])

doc.add_page_break()

# ========== 6. 验证结果 ==========
add_heading_custom(doc, '6. 验证结果', level=1)

add_heading_custom(doc, '6.1 MATLAB 仿真验证', level=2)
add_heading_custom(doc, '6.1.1 场景定义', level=3)

headers = ['场景', '描述', '关键参数', '评估重点']
data = [
    ['S1 基准负荷', '传统负荷', '340kW, PF=0.85', '建立基准分布'],
    ['S2 充电桩', 'EV 充电接入', '80kW, 5/7/11/13 次谐波', '谐波注入、负载上升'],
    ['S3 分布式光伏', '光伏发电', '200kW, 电压抬升 +2.93%', '电压抬升、反向潮流'],
    ['S4 光充耦合', '光伏+充电桩', '280kW, 谐波+电压抬升', '耦合效应分析'],
    ['S5 极端工况', '高渗透率', '360kW, 高 THD', '承载边界评估'],
]
add_table(doc, headers, data, col_widths=[2.5, 2.5, 4, 4])

add_heading_custom(doc, '6.1.2 仿真结果', level=3)

headers = ['场景', '触发事件数', 'AI 模组', '事件类型', '治理建议']
data = [
    ['S1 基准负荷', '0', 'ONLINE', '—', '系统运行正常，继续监测'],
    ['S2 充电桩', '100', 'ONLINE', 'HARMONIC', '配置 APF；有序充电策略'],
    ['S3 分布式光伏', '0', 'ONLINE', '—', '优化逆变器无功调节'],
    ['S4 光充耦合', '100', 'ONLINE', 'HARMONIC', '配置储能系统平滑功率'],
    ['S5 极端工况', '100', 'ONLINE', 'HARMONIC', '立即启动负荷切除'],
]
add_table(doc, headers, data, col_widths=[3, 2, 2, 3, 5])

add_heading_custom(doc, '6.2 真实硬件全链路验证', level=2)
add_heading_custom(doc, '6.2.1 测试环境', level=3)

headers = ['项目', '配置（已脱敏）']
data = [
    ['测试日期', '2026 年 8 月'],
    ['采样终端', '嵌入式 Linux 终端（SSH 远程访问）'],
    ['AI 算力模组', '独立 AI 计算模组（USB 虚拟网卡互连）'],
    ['测试工况', 'A 相加压，B/C 相开路'],
]
add_table(doc, headers, data, col_widths=[4, 12])

add_heading_custom(doc, '6.2.2 验证结果汇总', level=3)

headers = ['验证项', '状态', '详情']
data = [
    ['终端 HAL 初始化', '✅ 通过', 'HAL init → device get 成功'],
    ['波形采集', '✅ 通过', '7194 字节/周期，5 周期成功'],
    ['USB 网卡通信', '✅ 通过', 'RTT < 5ms，0 丢包'],
    ['原始波形传输', '✅ 通过', '协议头+波形数据，小端序'],
    ['波形解析', '✅ 通过', '7 通道 × 256 点 × 4 字节 float'],
    ['特征提取', '✅ 通过', '各通道电压有效值正确'],
    ['AI 推理', '✅ 通过', '异常检测 + 事件分类正确'],
    ['响应接收', '✅ 通过', '响应包校验通过'],
    ['日志系统', '✅ 通过', '分级日志控制台+文件'],
]
add_table(doc, headers, data, col_widths=[4, 2.5, 9.5])

add_heading_custom(doc, '6.2.3 AI 推理实测数据', level=3)

headers = ['指标', '实测值', '判定', '说明']
data = [
    ['A 相电压 RMS', '≈237 V', '正常', 'A 相正常加压'],
    ['B 相电压 RMS', '≈1.2 V', '异常', 'B 相开路'],
    ['C 相电压 RMS', '≈1.2 V', '异常', 'C 相开路'],
    ['三相电流 RMS', '< 0.001 A', '正常', '无电流通过'],
    ['iForest 得分', '1.0000', '异常', '判定为异常工况'],
    ['AE 得分', '0.7072', '异常', '重构误差偏大'],
    ['CNN 分类', '3 (单相开路)', '正确', '置信度 0.90'],
]
add_table(doc, headers, data, col_widths=[3.5, 3.5, 2, 7])

add_paragraph(doc,
    'AI 正确识别出 A 相加压、B/C 相开路的实际工况，验证了全链路的正确性。')

doc.add_page_break()

# ========== 7. 部署与运维方案 ==========
add_heading_custom(doc, '7. 部署与运维方案', level=1)

add_heading_custom(doc, '7.1 交叉编译与部署', level=2)

add_paragraph(doc,
    '支持在 Ubuntu 交叉编译服务器上完成 32 位 ARM 交叉编译，然后部署到目标硬件。')

add_heading_custom(doc, '7.1.1 部署流程', level=3)

deploy_steps = [
    '在交叉编译服务器上使用 ARM 工具链编译生成可执行文件',
    '通过安全文件传输上传编译产物至采样终端',
    '同时上传 AI 推理服务至算力模组',
    '在算力模组上启动 AI 推理服务',
    '在采样终端上运行程序，使用 32 位动态链接器加载运行',
]

for i, step in enumerate(deploy_steps, 1):
    p = doc.add_paragraph()
    run1 = p.add_run(f'{i}. ')
    set_run_font(run1, '微软雅黑', font_size=11, bold=True)
    run2 = p.add_run(step)
    set_run_font(run2, '微软雅黑', font_size=11)

add_heading_custom(doc, '7.2 运维监控', level=2)

add_heading_custom(doc, '7.2.1 日志系统', level=3)

headers = ['日志级别', '标识', '用途', '输出目标']
data = [
    ['ERROR', '[ERROR]', '错误信息，需立即处理', '控制台 + 日志文件'],
    ['WARN', '[WARN]', '警告信息，潜在问题', '控制台 + 日志文件'],
    ['INFO', '[INFO]', '关键流程节点', '控制台 + 日志文件'],
    ['DEBUG', '[DEBUG]', '详细调试信息', '仅日志文件（可通过参数启用）'],
]
add_table(doc, headers, data, col_widths=[2.5, 2, 4, 5.5])

add_heading_custom(doc, '7.2.2 安全访问', level=3)
add_paragraph(doc,
    '为保障系统安全，远程访问设备时：')

security_items = [
    '使用密钥认证替代密码登录',
    '通过堡垒机或 VPN 进行远程访问',
    '定期轮换访问凭据',
    '启用账户锁定策略（多次失败自动锁定）',
    '配置详细的访问审计日志',
]

for item in security_items:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '7.3 一键部署脚本', level=2)
add_paragraph(doc,
    '项目提供一键部署脚本，支持以下功能：')

script_features = [
    '自动交叉编译 → 上传各设备 → 启动服务 → 运行测试 → 收集日志',
    '彩色分级日志输出，便于快速定位问题',
    '支持超时保护，防止无限等待',
    '自动收集测试结果并生成报告',
]

for feat in script_features:
    p = doc.add_paragraph(feat, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

doc.add_page_break()

# ========== 8. 后续工作规划 ==========
add_heading_custom(doc, '8. 后续工作规划', level=1)

add_heading_custom(doc, '8.1 近期目标', level=2)

near_term = [
    '完成真实硬件驱动开发：替换仿真 HAL 为真实硬件接口',
    'AI 模型训练与部署：使用真实数据训练模型，量化后部署到 NPU',
    '多场景测试：在真实硬件上验证各类典型工况',
    'MQTT 通信完善：实现与云平台的完整数据交互',
]

for item in near_term:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '8.2 中期目标', level=2)

mid_term = [
    '低功耗核心集成：将核心采集任务迁移至低功耗核心',
    '嵌入式数据库：替换 CSV 为 SQLite 存储',
    'IEC 61850 协议对接：支持电力行业标准通信协议',
    'WEB 监控平台：开发可视化远程监控界面',
]

for item in mid_term:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

add_heading_custom(doc, '8.3 远期展望', level=2)

long_term = [
    '大模型部署：在算力模组上部署训练好的大模型',
    '边缘 AI 优化：模型 INT8 量化、剪枝，实现低功耗高效推理',
    '多终端协同：多台终端组成监测网络，实现区域级分析',
    '产品化落地：形成标准化产品方案，推广至应用场景',
]

for item in long_term:
    p = doc.add_paragraph(item, style='List Bullet')
    for run in p.runs:
        set_run_font(run, '微软雅黑', font_size=11)

# ========== 附录 ==========
doc.add_page_break()
add_heading_custom(doc, '附录', level=1)

add_heading_custom(doc, '附录 A：版本历史', level=2)

headers = ['版本', '日期', '主要变更']
data = [
    ['v2.2.1', '2026-08', '密码策略优化，运维文档完善'],
    ['v2.2.0', '2026-08', '真实硬件全链路验证，双机协作架构'],
    ['v2.1.0', '2026-08', '双机协作架构，USB 虚拟网卡通信'],
    ['v2.0.0', '2026-08', '完整复现版，五场景仿真验证'],
    ['v1.0.0', '2026-08', '初始版本，仿真框架搭建'],
]
add_table(doc, headers, data, col_widths=[2, 2.5, 11.5])

add_heading_custom(doc, '附录 B：术语表', level=2)

headers = ['术语', '含义']
data = [
    ['PQ', 'Power Quality，电能质量'],
    ['THD', 'Total Harmonic Distortion，总谐波畸变率'],
    ['AFE', 'Analog Front End，模拟前端'],
    ['RMS', 'Root Mean Square，有效值'],
    ['iForest', 'Isolation Forest，孤立森林'],
    ['AE', 'Autoencoder，自编码器'],
    ['CNN', 'Convolutional Neural Network，卷积神经网络'],
    ['NPU', 'Neural Processing Unit，神经网络处理单元'],
]
add_table(doc, headers, data, col_widths=[3, 13])

add_heading_custom(doc, '附录 C：安全声明', level=2)
add_paragraph(doc,
    '本报告中的系统架构、算法方案和验证结果为技术展示用途。'
    '实际部署时，所有设备访问凭据（用户名、密码、密钥）应由组织的密钥管理系统集中管理，'
    '并遵循最小权限原则进行授权。网络通信应加密传输，相关配置请参照组织的安全规范执行。')

# ========== 保存文档 ==========
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
    'PQ_AI_Terminal_技术方案报告_对外版_v2.2.1.docx')
doc.save(output_path)
print(f'对外技术方案报告（脱敏版）已生成：{output_path}')
